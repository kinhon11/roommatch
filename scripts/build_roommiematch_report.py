from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"C:\Users\Kinh\Downloads\KTMT\Kinh - BCTT RoommieMatch ReactJS.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def set_run_font(run, size=13, bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", style=None, align=None, bold=False, italic=False, size=13):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.0) if style is None and text else None
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=15 if level == 1 else 13, bold=True, color=(31, 78, 121))
    if not p.runs:
        r = p.add_run(text)
        set_run_font(r, size=15 if level == 1 else 13, bold=True, color=(31, 78, 121))
    else:
        p.runs[0].text = text
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=13)
    return p


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F6F8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line in code.strip().splitlines():
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9.5)
        p.add_run("\n")
    doc.add_paragraph()


def add_info_table(doc):
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Thông tin thực tập", "Thông tin thực tập"),
        ("Tên đơn vị thực tập", "Công ty TNHH Stramark Việt Nam"),
        ("Tên bộ phận thực tập", "Kỹ thuật phát triển phần mềm"),
        ("Các nhiệm vụ thực tập", "Nghiên cứu, phân tích và phát triển website sử dụng ReactJS"),
        ("Thời gian thực tập", "Từ ngày 01/09/2025 đến ngày 26/12/2025"),
        ("Người hướng dẫn thực tập", "Chị Phạm Thị Linh Chi"),
    ]
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, text, bold=i == 0 or j == 0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(13)

    add_paragraph(doc, "TRƯỜNG ĐẠI HỌC VINH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    add_paragraph(doc, "VIỆN KỸ THUẬT VÀ CÔNG NGHỆ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "BÁO CÁO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20)
    add_paragraph(doc, "THỰC TẬP CUỐI KHÓA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20)
    doc.add_paragraph()
    add_paragraph(doc, "VỊ TRÍ THỰC TẬP:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15)
    add_paragraph(doc, "THỰC TẬP SINH PHÁT TRIỂN WEBSITE SỬ DỤNG REACTJS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=15)
    doc.add_paragraph()

    cover_table = doc.add_table(rows=2, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_data = [("ĐVTT:", "Công ty TNHH Stramark Việt Nam"), ("SVTH:", "................................................")]
    for i, row in enumerate(cover_data):
        for j, text in enumerate(row):
            set_cell_text(cover_table.cell(i, j), text, bold=j == 0)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Nghệ An, 12/2025", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    doc.add_page_break()

    add_heading(doc, "LỜI CẢM ƠN", 1)
    for text in [
        "Em xin gửi lời cảm ơn chân thành đến Công ty TNHH Stramark Việt Nam vì đã tạo điều kiện để em được tham gia thực tập tại vị trí Thực tập sinh phát triển website sử dụng ReactJS trong thời gian vừa qua.",
        "Trong quá trình thực tập, em có cơ hội tìm hiểu quy trình làm việc tại doanh nghiệp công nghệ, đồng thời trực tiếp tham gia phân tích, thiết kế và xây dựng website RoommieMatch. Đây là dự án hỗ trợ tìm kiếm phòng trọ, quản lý bài đăng, đặt lịch xem phòng, yêu thích phòng, đặt cọc, trò chuyện và gợi ý phòng phù hợp bằng AI.",
        "Thông qua dự án, em được rèn luyện kỹ năng xây dựng giao diện bằng ReactJS, tổ chức route bằng React Router, kết nối API bằng Axios, quản lý dữ liệu với Supabase và phối hợp giữa frontend với backend Node.js/Express. Những kiến thức này giúp em hiểu rõ hơn cách một sản phẩm website thực tế được triển khai từ yêu cầu ban đầu đến khi hoàn thiện chức năng.",
        "Em cũng xin cảm ơn quý thầy cô Viện Kỹ thuật và Công nghệ, Trường Đại học Vinh đã trang bị nền tảng kiến thức chuyên môn, giúp em có đủ cơ sở để tiếp cận công việc thực tế tại doanh nghiệp.",
        "Do thời gian thực tập và kinh nghiệm thực tế còn hạn chế, báo cáo không tránh khỏi những thiếu sót. Em kính mong nhận được sự góp ý của quý thầy cô để có thể hoàn thiện hơn trong học tập và công việc sau này.",
        "Em xin chân thành cảm ơn!",
    ]:
        add_paragraph(doc, text)
    add_paragraph(doc, "Nghệ An, tháng 12 năm 2025", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "Sinh viên", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "................................................", align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    add_heading(doc, "GIỚI THIỆU ĐƠN VỊ THỰC TẬP", 1)
    add_heading(doc, "1.1. Giới thiệu Công ty TNHH Stramark Việt Nam", 2)
    add_heading(doc, "1.1.1. Tổng quan về công ty", 2)
    add_paragraph(doc, "Công ty TNHH Stramark Việt Nam là doanh nghiệp hoạt động trong lĩnh vực công nghệ thông tin và phát triển phần mềm, tập trung cung cấp các giải pháp website, hệ thống quản lý và dịch vụ công nghệ cho doanh nghiệp. Công ty chú trọng vào chất lượng sản phẩm, tính ổn định, khả năng mở rộng và trải nghiệm người dùng trong quá trình triển khai dự án.")
    add_paragraph(doc, "Môi trường làm việc tại công ty tạo điều kiện cho sinh viên thực tập tiếp cận quy trình phát triển phần mềm thực tế, từ phân tích yêu cầu, thiết kế giao diện, xây dựng chức năng, kiểm thử đến chỉnh sửa và hoàn thiện sản phẩm theo phản hồi.")
    add_heading(doc, "Giá trị cốt lõi", 2)
    for item in [
        "Khách hàng là trọng tâm: luôn đặt nhu cầu sử dụng thực tế và hiệu quả vận hành của khách hàng làm ưu tiên.",
        "Chất lượng và hiệu quả: hướng tới sản phẩm ổn định, dễ sử dụng, dễ bảo trì và có giá trị ứng dụng.",
        "Không ngừng học hỏi và đổi mới: khuyến khích nghiên cứu, áp dụng công nghệ mới trong phát triển phần mềm.",
        "Tinh thần trách nhiệm: đề cao sự chủ động, minh bạch, đúng tiến độ và phối hợp hiệu quả trong công việc.",
        "Phát triển bền vững: xây dựng giải pháp có khả năng mở rộng, phù hợp với định hướng lâu dài của doanh nghiệp.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.1.2. Tầm nhìn và sứ mệnh", 2)
    add_paragraph(doc, "Stramark Việt Nam hướng tới trở thành đơn vị phát triển phần mềm và website uy tín, cung cấp các giải pháp công nghệ phù hợp với nhu cầu số hóa của doanh nghiệp. Sứ mệnh của công ty là hỗ trợ khách hàng xây dựng, vận hành và tối ưu các nền tảng số thông qua quy trình phát triển chuyên nghiệp và công nghệ hiện đại.")
    add_heading(doc, "1.1.3. Lĩnh vực kinh doanh và dịch vụ", 2)
    for item in [
        "Phát triển website: thiết kế, xây dựng và tối ưu website hiện đại trên nền tảng ReactJS, VueJS hoặc các framework phù hợp.",
        "Phát triển phần mềm theo yêu cầu: xây dựng hệ thống quản lý, cổng thông tin và ứng dụng web theo nghiệp vụ riêng.",
        "Giải pháp công nghệ cho doanh nghiệp: tư vấn số hóa quy trình, tích hợp API, lưu trữ dữ liệu và triển khai hệ thống.",
        "Nghiên cứu và ứng dụng công nghệ mới: áp dụng các công nghệ frontend, backend, cơ sở dữ liệu và AI nhằm nâng cao hiệu quả sản phẩm.",
    ]:
        add_bullet(doc, item)
    add_info_table(doc)

    add_heading(doc, "KẾT QUẢ THỰC TẬP", 1)
    add_heading(doc, "2.1. Cơ sở lý thuyết và công nghệ nền tảng", 2)
    add_paragraph(doc, "Trong thời gian thực tập tại Công ty TNHH Stramark Việt Nam, em được định hướng nghiên cứu và xây dựng website RoommieMatch sử dụng ReactJS. Nội dung thực tập tập trung vào xây dựng giao diện người dùng, tổ chức tuyến đường, xử lý trạng thái đăng nhập, kết nối API, tích hợp Supabase và hoàn thiện các chức năng phục vụ người thuê, chủ trọ, môi giới và quản trị viên.")
    add_heading(doc, "2.1.1. ReactJS và kiến trúc component", 2)
    add_paragraph(doc, "ReactJS là thư viện JavaScript dùng để xây dựng giao diện người dùng theo kiến trúc component. Mỗi thành phần giao diện được tách thành một component độc lập, có thể tái sử dụng và kết hợp với nhau để tạo thành trang hoàn chỉnh. Cách tổ chức này giúp mã nguồn rõ ràng, dễ bảo trì và thuận tiện khi mở rộng chức năng.")
    add_paragraph(doc, "Trong RoommieMatch, các trang được tổ chức theo nhóm vai trò như Tenant, Landlord, Broker, Admin và AI. Các component dùng chung như Navbar, ProtectedRoute, NotificationBell và AIChatWidget được tách riêng để giảm lặp mã và tăng tính nhất quán của giao diện.")
    add_code(doc, """
import { Suspense, lazy } from 'react';
import { Routes, Route } from 'react-router-dom';

const RoomsPage = lazy(() => import('./pages/Tenant/RoomsPage'));
const RoomDetail = lazy(() => import('./pages/Tenant/RoomDetail'));
const AdminDashboard = lazy(() => import('./pages/Admin/Dashboard'));
""")
    add_heading(doc, "2.1.2. Vite và môi trường phát triển frontend", 2)
    add_paragraph(doc, "Dự án sử dụng Vite làm công cụ khởi tạo và build frontend. Vite hỗ trợ thời gian khởi động nhanh, hot module replacement và cấu hình gọn nhẹ, phù hợp với quá trình phát triển ứng dụng React hiện đại.")
    add_heading(doc, "2.1.3. React Router và điều hướng theo vai trò", 2)
    add_paragraph(doc, "React Router được sử dụng để quản lý điều hướng trong ứng dụng một trang. Các route công khai như trang chủ, danh sách phòng, chi tiết phòng được tách biệt với các route yêu cầu đăng nhập như hồ sơ, lịch hẹn, tin nhắn, đặt cọc và trang quản trị.")
    add_code(doc, """
<Route
  path="/profile"
  element={
    <ProtectedRoute allowedRoles={['tenant', 'landlord', 'broker', 'admin']}>
      <ProfilePage />
    </ProtectedRoute>
  }
/>
""")
    add_heading(doc, "2.1.4. Quản lý xác thực và phân quyền", 2)
    add_paragraph(doc, "Dự án sử dụng cơ chế xác thực người dùng thông qua backend và lưu thông tin phiên đăng nhập ở phía client. AuthContext cung cấp trạng thái đăng nhập, thông tin người dùng và các hàm login, register, logout để các component có thể sử dụng thống nhất.")
    add_code(doc, """
const login = useCallback(async (credentials) => {
  const response = await authService.login(credentials);
  setUser(response.user);
  setSession(response.session);
  localStorage.setItem('roommie-user', JSON.stringify(response.user));
  return response;
}, []);
""")
    add_heading(doc, "2.1.5. Backend Node.js/Express và Supabase", 2)
    add_paragraph(doc, "Bên cạnh frontend ReactJS, RoommieMatch sử dụng backend Node.js/Express để xây dựng API, xử lý nghiệp vụ và giao tiếp với Supabase. Supabase PostgreSQL được dùng làm cơ sở dữ liệu chính, đồng thời hỗ trợ lưu trữ dữ liệu người dùng, phòng trọ, lịch hẹn, yêu thích, đặt cọc, tin nhắn và thông báo.")

    add_heading(doc, "2.2. Xây dựng và phát triển website RoommieMatch", 2)
    add_heading(doc, "2.2.1. Giới thiệu dự án RoommieMatch", 2)
    add_paragraph(doc, "RoommieMatch là website hỗ trợ tìm kiếm phòng trọ và kết nối các nhóm người dùng liên quan đến nhu cầu thuê phòng. Hệ thống cho phép người thuê tra cứu phòng, xem chi tiết, lưu yêu thích, đặt lịch xem phòng, gửi yêu cầu ghép phòng, đặt cọc và trò chuyện. Chủ trọ có thể đăng tin, quản lý phòng và xử lý yêu cầu. Quản trị viên có thể duyệt tin, quản lý người dùng và theo dõi báo cáo.")
    add_paragraph(doc, "Mục tiêu của dự án là xây dựng một nền tảng web có giao diện dễ sử dụng, hỗ trợ nhiều vai trò người dùng, đảm bảo dữ liệu được quản lý tập trung và có khả năng mở rộng thêm các tính năng thông minh như gợi ý phòng bằng AI.")
    add_heading(doc, "Công nghệ sử dụng", 2)
    for item in [
        "Frontend: ReactJS, Vite, React Router, Axios, Supabase JS.",
        "Backend: Node.js, Express, JWT, Multer, Supabase JS.",
        "Database: Supabase PostgreSQL.",
        "AI: Google Gemini và MiniMax AI cho chức năng trợ lý/gợi ý phòng.",
        "Công cụ kiểm tra: ESLint, Node Test Runner và quy trình build frontend.",
    ]:
        add_bullet(doc, item)
    add_code(doc, """
\"dependencies\": {
  \"react\": \"^19.2.4\",
  \"react-dom\": \"^19.2.4\",
  \"react-router-dom\": \"^7.13.2\",
  \"axios\": \"^1.16.0\",
  \"@supabase/supabase-js\": \"^2.104.1\"
}
""")
    add_heading(doc, "2.2.2. Triển khai các chức năng chính", 2)
    add_heading(doc, "2.2.2.1. Đăng ký, đăng nhập và phân quyền người dùng", 2)
    add_paragraph(doc, "Hệ thống hỗ trợ người dùng đăng ký, đăng nhập và truy cập chức năng theo vai trò. Các vai trò chính gồm tenant, landlord, broker và admin. ProtectedRoute kiểm tra trạng thái đăng nhập và quyền truy cập trước khi hiển thị trang chức năng, giúp hạn chế người dùng truy cập nhầm khu vực không thuộc quyền của mình.")
    add_heading(doc, "2.2.2.2. Tìm kiếm và xem chi tiết phòng trọ", 2)
    add_paragraph(doc, "Người thuê có thể xem danh sách phòng đã được duyệt, lọc theo nhu cầu và mở trang chi tiết để xem thông tin phòng. Dữ liệu phòng được lấy từ API thông qua roomService, giúp tách riêng phần giao diện với phần gọi dữ liệu.")
    add_code(doc, """
getApprovedRooms: async (filters = {}) => {
  const { data } = await apiClient.get('/rooms', { params: filters });
  return data;
},

getRoomById: async (id) => {
  const { data } = await apiClient.get(`/rooms/${id}`);
  return data;
}
""")
    add_heading(doc, "2.2.2.3. Quản lý bài đăng phòng cho chủ trọ", 2)
    add_paragraph(doc, "Đối với chủ trọ, hệ thống cung cấp các chức năng đăng phòng mới, chỉnh sửa thông tin phòng, tải ảnh phòng, ẩn/hiện phòng và cập nhật trạng thái còn phòng. Các chức năng này giúp chủ trọ chủ động quản lý dữ liệu phòng trọ của mình trên website.")
    add_heading(doc, "2.2.2.4. Lịch hẹn, yêu thích, đặt cọc và trò chuyện", 2)
    add_paragraph(doc, "RoommieMatch hỗ trợ các thao tác quan trọng trong quá trình thuê phòng như lưu phòng yêu thích, đặt lịch xem phòng, gửi yêu cầu, đặt cọc và trò chuyện giữa người dùng. Các chức năng này giúp quy trình tìm và thuê phòng diễn ra thuận tiện hơn, giảm việc trao đổi rời rạc qua nhiều kênh khác nhau.")
    add_heading(doc, "2.2.2.5. Trang quản trị hệ thống", 2)
    add_paragraph(doc, "Quản trị viên có quyền truy cập các trang quản lý người dùng, duyệt phòng chờ, quản lý toàn bộ phòng và xử lý báo cáo. Việc tách riêng khu vực admin giúp hệ thống kiểm soát chất lượng tin đăng, hạn chế nội dung không phù hợp và nâng cao độ tin cậy của nền tảng.")
    add_heading(doc, "2.2.2.6. Trợ lý AI gợi ý phòng", 2)
    add_paragraph(doc, "Dự án tích hợp trợ lý AI nhằm hỗ trợ người dùng đặt câu hỏi, mô tả nhu cầu thuê phòng và nhận gợi ý phù hợp. Chức năng này được tổ chức thành các component riêng như AIChatWidget, AIMessageList, AIQuickPrompts và các service backend xử lý prompt, ngữ cảnh và điều phối công cụ.")
    add_paragraph(doc, "Việc tích hợp AI giúp website không chỉ dừng ở chức năng tra cứu thông thường mà còn hỗ trợ người dùng ra quyết định nhanh hơn dựa trên nhu cầu, ngân sách và vị trí mong muốn.")
    add_heading(doc, "2.2.3. Kết quả đạt được", 2)
    for item in [
        "Hoàn thiện cấu trúc frontend ReactJS với các trang chức năng theo từng vai trò người dùng.",
        "Kết nối frontend với backend thông qua hệ thống service và API client.",
        "Xây dựng luồng đăng nhập, phân quyền và bảo vệ route.",
        "Triển khai các chức năng cốt lõi: tìm phòng, chi tiết phòng, yêu thích, lịch hẹn, đặt cọc, trò chuyện, quản lý phòng và quản trị.",
        "Tìm hiểu cách tích hợp dịch vụ AI vào website để hỗ trợ gợi ý và tư vấn cho người dùng.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "BÀI HỌC KINH NGHIỆM", 1)
    add_heading(doc, "3.1. Bài học về vận dụng kiến thức đã học", 2)
    add_paragraph(doc, "Qua quá trình thực tập, em nhận thấy kiến thức nền tảng về lập trình web, cơ sở dữ liệu, mạng máy tính và quy trình phát triển phần mềm có vai trò rất quan trọng khi triển khai dự án thực tế. Những kiến thức về HTTP, RESTful API, JSON, xác thực người dùng và tổ chức giao diện giúp em hiểu rõ hơn cách các thành phần trong một website phối hợp với nhau.")
    add_paragraph(doc, "Việc xây dựng RoommieMatch cũng giúp em vận dụng tư duy phân tích bài toán, chia nhỏ chức năng và tổ chức mã nguồn theo từng nhóm nhiệm vụ. Thay vì chỉ viết giao diện tĩnh, em học được cách xử lý dữ liệu động, trạng thái đăng nhập, quyền truy cập và các tình huống phát sinh trong quá trình người dùng thao tác.")
    add_heading(doc, "3.2. Bài học từ đơn vị thực tập", 2)
    add_paragraph(doc, "Môi trường thực tập tại Công ty TNHH Stramark Việt Nam giúp em tiếp cận quy trình làm việc chuyên nghiệp hơn, đặc biệt là cách đọc yêu cầu, phân tích chức năng, trao đổi khi gặp vướng mắc và hoàn thiện sản phẩm theo từng giai đoạn. Em hiểu rằng một sản phẩm phần mềm không chỉ cần chạy được mà còn cần dễ sử dụng, dễ bảo trì và phù hợp với nhu cầu thực tế.")
    add_paragraph(doc, "Thông qua dự án, em học thêm được cách tổ chức source code ReactJS, tách component dùng chung, xây dựng service gọi API, kết hợp frontend với backend và kiểm tra lỗi trong quá trình phát triển. Đây là những kinh nghiệm quan trọng giúp em tự tin hơn khi tham gia các dự án web sau này.")
    add_heading(doc, "3.3. Bài học về kỹ năng và phẩm chất nghề nghiệp", 2)
    for item in [
        "Kỹ năng tự học: chủ động đọc tài liệu, tìm hiểu thư viện và thử nghiệm giải pháp khi gặp yêu cầu mới.",
        "Kỹ năng giải quyết vấn đề: phân tích lỗi theo từng bước, xác định nguyên nhân và kiểm tra lại sau khi sửa.",
        "Tư duy mã sạch: đặt tên biến, tách file và tổ chức component rõ ràng để mã nguồn dễ đọc, dễ bảo trì.",
        "Tinh thần trách nhiệm: hoàn thành nhiệm vụ đúng tiến độ, cẩn thận khi chỉnh sửa chức năng ảnh hưởng đến người dùng.",
        "Kỹ năng giao tiếp: trao đổi rõ ràng với người hướng dẫn và tiếp nhận góp ý để cải thiện sản phẩm.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.4. Kiến nghị với Nhà trường", 2)
    add_paragraph(doc, "Từ trải nghiệm thực tế tại doanh nghiệp, em xin đề xuất Nhà trường tiếp tục tăng cường các học phần thực hành theo dự án, đặc biệt là các công nghệ phát triển website hiện đại như ReactJS, Node.js, Express và cơ sở dữ liệu PostgreSQL/Supabase. Việc sinh viên được làm quen sớm với quy trình xây dựng sản phẩm thực tế sẽ giúp quá trình thực tập và làm việc sau khi tốt nghiệp thuận lợi hơn.")
    add_paragraph(doc, "Bên cạnh đó, Nhà trường có thể bổ sung thêm nội dung về Git, làm việc nhóm, kiểm thử phần mềm, triển khai ứng dụng và bảo mật cơ bản trong phát triển web. Đây là những kỹ năng thường xuyên được sử dụng trong môi trường doanh nghiệp và có ý nghĩa thiết thực đối với sinh viên ngành công nghệ thông tin.")

    add_heading(doc, "KẾT LUẬN", 1)
    add_paragraph(doc, "Kỳ thực tập tại Công ty TNHH Stramark Việt Nam là cơ hội quan trọng giúp em củng cố kiến thức đã học và tiếp cận công việc phát triển website trong môi trường thực tế. Thông qua dự án RoommieMatch, em được rèn luyện kỹ năng xây dựng frontend bằng ReactJS, kết nối backend, quản lý dữ liệu, xử lý xác thực người dùng và triển khai các chức năng phục vụ nhiều nhóm người dùng khác nhau.")
    add_paragraph(doc, "Những kinh nghiệm thu được trong thời gian thực tập là nền tảng để em tiếp tục hoàn thiện kỹ năng chuyên môn, nâng cao tinh thần trách nhiệm và định hướng phát triển nghề nghiệp trong lĩnh vực lập trình website.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
